from docx import Document
from docx.oxml import parse_xml, OxmlElement, ns
from docx.shared import Pt, Cm
from tasks import task1, task2, task3, task4, task5, task6, task7, task8, task9, task10, task11, task12, task13, task14, task15, task16, task17, task18

def generateDocument(taskOutputFile="docx/tasks.docx", answerOutputFile="docx/answers.docx", variants=1, tasks=range(1, 19)):
    # Создание документов для заданий и ответов
    task_doc = Document()
    set_doc_margins(task_doc, top=0.5, right=0.5, bottom=0.5, left=0.5)
    answer_doc = Document()
    set_doc_margins(answer_doc, top=0.5, right=0.5, bottom=0.5, left=0.5)
    
    # Генерация вариантов
    for variant in range(1, variants+1):
        task_doc.add_heading(f"Вариант {variant}", level=1)
        answer_doc.add_heading(f"Вариант {variant}", level=1)
        
        # Генерация задач для текущего варианта
        for i in range(len(tasks)):
            task_number = tasks[i]
            task_module = globals()[f"task{task_number}"]
            function_name = f"generate_task{task_number}"
            task_function = getattr(task_module, function_name)
            task_function(task_doc, answer_doc, i+1)
        
        # Добавление разделителя между вариантами
        task_doc.add_page_break()
        answer_doc.add_page_break()
    
    # Сохранение документов
    task_doc.save(taskOutputFile)
    answer_doc.save(answerOutputFile)

def set_doc_margins(doc, top, right, bottom, left):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
