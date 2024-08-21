import random
import math
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            for attr, value in kwargs[edge].items():
                edge_element.set(qn(f'w:{attr}'), str(value))
            tcBorders.append(edge_element)


def generate_task14(task_doc, answer_doc, task_number):

    p = random.randint(1, 10)/10000
    q = 1-p
    n = random.randint(5, 15)*100

    task_text = f"Торговая база получила {n} электрических лампочек. Вероятность повреждения электролампочки в пути равна {p}. Составить ряд распределения числа лампочек, поврежденных в пути. Найти M(X) этой случайной величины."

    answer_doc.add_paragraph(f"{task_number})")

    table = answer_doc.add_table(rows=2, cols=7)
    table.cell(0, 0).text = 'x'
    table.cell(1, 0).text = 'p'
    for i in range(1, 7):
        if (i == 3 or i == 5):
            table.cell(0, i).text = "..."
            table.cell(1, i).text = "..."
        elif (i == 4):
            table.cell(0, i).text = "k"
            table.cell(1, i).text = str(f"C({n}, k) * {p}^k * {q}^({n}-k)")
        else:
            table.cell(0, i).text = str(i-1)
            if (i == 1):
                table.cell(1, i).text = str(f"{q}^{n}")
            if (i == 2):
                table.cell(1, i).text = str(f"{n}*{p}*{q}^{n-1}")
            if (i == 6):
                table.cell(1, i).text = str(f"{p}^{n}")
   
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

    MX = round(n*p, 4)

    task_answer = f"M(X) = {MX}"
    
    task_doc.add_paragraph(f"{task_number}) {task_text}")
    answer_doc.add_paragraph(f"{task_answer}")