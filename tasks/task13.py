import random
import math
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            for attr, value in kwargs[edge].items():
                edge_element.set(qn(f'w:{attr}'), str(value))
            tcBorders.append(edge_element)


def generate_task13(task_doc, answer_doc, task_number):

    p = q = 0.5
    n = random.randint(3, 5)

    task_text = f"Предполагая одинаковой вероятность рождения мальчика и девочки, составить ряд распределения случайной величины X, которая выражает число мальчиков в семье, имеющей {n} детей. Найти M(X) и D(X) этой случайной величины."

    answer_doc.add_paragraph(f"{task_number})")

    table = answer_doc.add_table(rows=2, cols=n+2)
    table.cell(0, 0).text = 'x'
    table.cell(1, 0).text = 'p'
    for i in range(1, n+2):
        table.cell(0, i).text = str(i-1)
        table.cell(1, i).text = str(f"1/{n+1}")

    # Добавление границ к таблице
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

    MX = round(n/2, 4)
    DX = round(((n+1)**2 - 1)/12, 4)

    task_answer = f"M(X) = {MX}; D(X) = {DX}"
    
    task_doc.add_paragraph(f"{task_number}) {task_text}")
    answer_doc.add_paragraph(f"{task_answer}")