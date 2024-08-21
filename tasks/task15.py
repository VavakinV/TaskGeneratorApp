import random
import math
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            for attr, value in kwargs[edge].items():
                edge_element.set(qn(f'w:{attr}'), str(value))
            tcBorders.append(edge_element)


def generate_task15(task_doc, answer_doc, task_number):

    task_text = f"Независимые случайные величины X и Y заданы таблицами распределений.\nНайти:\n1) M(X), M(Y), D(X), D(Y);\n2) таблицы распределения случайных величин Z = 2X + Y, W = X ⋅ Y;\n3) M(Z), M(W), D(Z), D(W) непосредственно по таблицам распределений и на основании свойств математического ожидания и дисперсии."

    x_arr = [0, 0, 0]
    px_arr = [0, 0, 0]
    y_arr = [0, 0]
    py_arr = [0, 0]

    x_arr[0] = random.randint(-3, 1)
    x_arr[1] = x_arr[0]+4
    x_arr[2] = x_arr[0]+6

    px_arr[0] = random.randint(1, 3)/10
    px_arr[1] = random.randint(1, 4)/10
    px_arr[2] = (10 - px_arr[0]*10 - px_arr[1]*10)/10

    y_arr[0] = random.randint(-4, -1)
    y_arr[1] = y_arr[0]+5

    py_arr[0] = random.randint(1, 5)/10
    py_arr[1] = (10 - py_arr[0]*10)/10

    task_doc.add_paragraph(f"{task_number}) {task_text}")

    tableX = task_doc.add_table(rows=2, cols=4)

    tableX.cell(0, 0).text = 'X'
    tableX.cell(1, 0).text = 'p'

    for i in range(1, 4):
        tableX.cell(0, i).text = str(x_arr[i-1])
        tableX.cell(1, i).text = str(px_arr[i-1])

    for row in tableX.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            ) 

    task_doc.add_paragraph()

    tableY = task_doc.add_table(rows=2, cols=3)

    tableY.cell(0, 0).text = 'Y'
    tableY.cell(1, 0).text = 'p'

    for i in range(1, 3):
        tableY.cell(0, i).text = str(y_arr[i-1])
        tableY.cell(1, i).text = str(py_arr[i-1])

    for row in tableY.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            ) 

    MX = 0
    for i in range(3):
        MX += x_arr[i]*px_arr[i]

    DX = 0
    for i in range(3):
        DX += px_arr[i] * ((x_arr[i] - MX)**2)
    
    MY = 0
    for i in range(2):
        MY += y_arr[i]*py_arr[i]
    
    DY = 0
    for i in range(2):
        DY += py_arr[i] * ((y_arr[i] - MY)**2)

    answer_doc.add_paragraph(f"{task_number}) M(X) = {round(MX, 4)}; M(Y) = {round(MY, 4)}; D(X) = {round(DX, 4)}; D(Y) = {round(DY, 4)}")

    z_arr = {}
    for i in range(3):
        for j in range(2):
            value = 2*x_arr[i] + y_arr[j]
            if str(value) in z_arr.keys():
                z_arr[str(value)] += px_arr[i]*py_arr[j]
            else:
                z_arr[str(value)] = round(px_arr[i]*py_arr[j], 4)
    
    z_arr = dict(sorted(z_arr.items(), key=lambda x:int(x[0])))
    l = len(z_arr.items())

    MZ = 0
    for i in range(l):
        MZ += z_arr[list(z_arr.keys())[i]] * int(list(z_arr.keys())[i])
    
    DZ = 0
    for i in range(l):
        DZ += z_arr[list(z_arr.keys())[i]] * (((int(list(z_arr.keys())[i])) - MZ)**2)

    tableZ = answer_doc.add_table(rows=2, cols=l+1)
    tableZ.cell(0, 0).text = 'z'
    tableZ.cell(1, 0).text = 'p'
    for i in range(1, l+1):
        tableZ.cell(0, i).text = list(z_arr.keys())[i-1]
        tableZ.cell(1, i).text = str(z_arr[list(z_arr.keys())[i-1]])

    for row in tableZ.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )  

    answer_doc.add_paragraph()           

    w_arr = {}
    for i in range(3):
        for j in range(2):
            value = x_arr[i] * y_arr[j]
            if str(value) in w_arr.keys():
                w_arr[str(value)] += px_arr[i]*py_arr[j]
            else:
                w_arr[str(value)] = round(px_arr[i]*py_arr[j], 4)
    
    w_arr = dict(sorted(w_arr.items(), key=lambda x:int(x[0])))
    l = len(w_arr.items())

    MW = 0
    for i in range(l):
        MW += w_arr[list(w_arr.keys())[i]] * int(list(w_arr.keys())[i])
    
    DW = 0
    for i in range(l):
        DW += w_arr[list(w_arr.keys())[i]] * (((int(list(w_arr.keys())[i])) - MW)**2)

    tableW = answer_doc.add_table(rows=2, cols=l+1)
    tableW.cell(0, 0).text = 'w'
    tableW.cell(1, 0).text = 'p'
    for i in range(1, l+1):
        tableW.cell(0, i).text = list(w_arr.keys())[i-1]
        tableW.cell(1, i).text = str(w_arr[list(w_arr.keys())[i-1]])

    for row in tableW.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            ) 

    answer_doc.add_paragraph(f"M(Z) = {round(MZ, 4)}; M(W) = {round(MW, 4)}; D(Z) = {round(DZ, 4)}; D(W) = {round(DW, 4)}")