import random
import math
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            for attr, value in kwargs[edge].items():
                edge_element.set(qn(f'w:{attr}'), str(value))
            tcBorders.append(edge_element)


def generate_task12(task_doc, answer_doc, task_number):

    p = random.randint(1, 5)/10
    q = (1-p)

    task_text = f"Вероятность изготовления нестандартной детали равна {p}. Из партии контролер проверяет не более четырех деталей. Если деталь оказывается нестандартной, испытания прекращаются, а партия задерживается. Если деталь оказывается стандартной, контролер берет следующую и т. д. Составить ряд распределения числа проверенных деталей. Найти М(Х), D(X), σ(X), F(X) этой случайной величины. Построить график F(X)."

    p*=10
    q*=10

    ps = [0, 0, 0, 0]
    ps[0] = p/10
    ps[1] = q*p/100
    ps[2] = q*q*p/1000
    ps[3] = q*q*q*p/10000

    p/=10
    q/=10

    answer_doc.add_paragraph(f"{task_number})")

    table = answer_doc.add_table(rows=2, cols=5)
    table.cell(0, 0).text = 'x'
    table.cell(1, 0).text = 'p'
    for i in range(1, 5):
        table.cell(0, i).text = str(i)
        table.cell(1, i).text = str(ps[i-1])

    # Добавление границ к таблице
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

    MX = round(1/p, 4)
    DX = round(q*10/(p*10)**2 * 10, 4)
    sigmaX = round(math.sqrt(DX), 4)

    task_answer = f"M(X) = {MX}; D(X) = {DX}; σ(X) = {sigmaX}; F(x) = 1 - {q}"
    
    task_doc.add_paragraph(f"{task_number}) {task_text}")

    paragraph = answer_doc.add_paragraph(f"{task_answer}")

    run = paragraph.add_run('x-1')
    run.font.superscript = True